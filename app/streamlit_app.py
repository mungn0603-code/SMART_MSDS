"""서비스 UI: 물질 선택 -> CAMEO 판정 확정 -> MSDS 근거 직접조회(§2/§10) -> LLM 설명.

파이프라인 구성은 run_cameo_full.py(--context pair)와 동일하게 맞춘다 —
프롬프트(run_cameo_context_pilot.SYSTEM_PROMPT/build_prompt), CAMEO 컨텍스트
(cameo_group_lookup.format_context(detailed=True)), 근거 수집 규칙(corpus_tag='service'의
§2·§10을 두 CAS로 직접 조회하고 정렬을 고정), LLM 파라미터(generate_baseline 상수)를
전부 그대로 재사용한다. 여기서 새로 하는 건 UI뿐이다.

**근거 수집에 검색을 쓰지 않는다**(2026-08-29). 사용자가 목록에서 고르므로 CAS가 이미
정해져 있고, 정답 근거는 정의상 그 두 물질의 §2다 - 사유와 실측은 msds_context()의
docstring과 docs/RETRIEVAL.md 5절. retrieve()는 자유 문장 질의를 위해 남겨 둔다.

UI는 st.set_page_config 이후 커스텀 CSS(_inject_css)로 데모/AI-생성 느낌을 지우고
"현업 EHS 대시보드"에 가까운 톤(중립 배경 + 단일 accent + 뱃지형 판정 + 히트맵
매트릭스)으로 다시 그렸다. 판정 로직/검색/프롬프트는 손대지 않았다 — 프레젠테이션
레이어(main() 하단부)만 재작성.

  streamlit run app/streamlit_app.py
  python app/streamlit_app.py --check   # LLM 없이 검색+CAMEO 경로만 자체점검
"""

from __future__ import annotations

import io
import itertools
import re
import sqlite3
import sys
from pathlib import Path

import numpy as np
import pandas as pd
import streamlit as st

ROOT = Path(__file__).resolve().parent.parent
sys.path[:0] = [str(ROOT / "src"), str(ROOT / "scripts" / "5_generation")]

import cameo_group_lookup as CL  # noqa: E402
import generate_baseline as GB  # noqa: E402
import kr_glossary as KRG  # noqa: E402
import llm as L  # noqa: E402
import retrieval as R  # noqa: E402
from compatibility_engine import DISCLAIMER, CompatibilityEngine  # noqa: E402
from run_cameo_context_pilot import PROMPT_VERSION, build_prompt  # noqa: E402

DB_PATH = ROOT / "data" / "reactivity_reference.db"

# certifi 번들에 로컬 루트 CA가 없어 huggingface.co 인증서 검증이 실패하는 환경이다
# (llm.py와 같은 문제/같은 해법 — SSL_CERT_FILE로 Windows 저장소 번들 지정, 검증은 끄지 않음).
# 검증 실패 시 huggingface_hub이 재시도를 소진한 뒤 httpx 클라이언트를 닫아버려서 최종 오류가
# "Cannot send a request, as the client has been closed"로 둔갑한다 — SSL 문제라는 게 표면에
# 안 드러나니 여기 남긴다. HF를 건드리는 경로가 embedder()와 embed_corpus() 둘이라
# 함수 안이 아니라 import 시점에 한 번 건다(둘 중 어느 쪽이 먼저 돌든 적용되도록).
L._ensure_ca_bundle()

# freeze_retrieval.py와 동일한 baseline 구성. 바꾸면 README 실측치와 어긋난다.
MODEL, GRAN = "bge-m3-ko", "section"
SECTIONS = {2, 10}
CAND_K, TOPK = 20, 10

# 2026-08-28: 서비스 검색 대상은 corpus_tag='service' 하나다(seed_service_corpus.py).
# 이전에는 ('173','core') 합본을 썼는데, 그러면 Registry 밖 legacy 89종이 검색 후보로
# 남아 서비스 불가 물질이 retrieval 경쟁에 끼어든다(아세트산 <- '초산 에틸' 사례).
# '173'/'core' 태그는 평가 재현용으로 DB에 그대로 보존하되 서비스 경로에서 쓰지 않는다.
# 물질 수는 여기에 적지 않는다 - substance_status VIEW에서 계산된다.
SERVICE_CORPUS_TAG = "service"
CORPUS_TAGS = (SERVICE_CORPUS_TAG,)
BM25_TAG = "section_s210_" + SERVICE_CORPUS_TAG  # 코퍼스가 바뀌면 문서통계도 달라 BM25 재빌드

QUERY_TEMPLATE = "{a}, {b} 두 물질을 함께 취급해도 되는가? 혼합 시 위험성과 유의사항은?"

# msds_sections.section -> 화면 표시용 라벨(KOSHA MSDS 항목 번호 그대로)
SECTION_LABEL = {
    2: "§2 유해성·위험성",
    3: "§3 구성성분",
    9: "§9 물리화학적 특성",
    10: "§10 안정성 및 반응성",
}

CATEGORY_LABEL = {
    "Incompatible": "부적합 · Incompatible",
    "Caution": "주의 · Caution",
    "Compatible": "적합 · Compatible",
    "Abstain": "판단 보류 · Abstain",
}

CELL_STYLE = {
    "Incompatible": "background-color:#FEF2F2;color:#B91C1C;font-weight:600;",
    "Caution": "background-color:#FFFBEB;color:#92400E;font-weight:600;",
    "Compatible": "background-color:#F0FDF4;color:#15803D;font-weight:600;",
    "Abstain": "background-color:#F8FAFC;color:#64748B;font-weight:600;",
    "—": "color:#D1D5DB;",
}

# PAIR INSPECTOR 목록 앞에 붙는 판정 색점. Streamlit 버튼은 HTML을 못 받으므로
# 유니코드 원으로 대체(디자인 시안이 이미 이 표기를 사용).
SEVERITY_DOT = {"Incompatible": "🔴", "Caution": "🟡", "Compatible": "🟢", "Abstain": "⚪"}

# 최종 보고서 표의 "판정" 셀 배지 색(hex, # 없이) - Streamlit(mv-badge)과 DOCX/PDF export가
# 같은 팔레트를 쓰도록 여기 한 곳에서만 정의한다(CELL_STYLE과 톤은 맞추되 폰트 색 대비를
# hex로 직접 써야 하는 reportlab/python-docx 때문에 별도 상수로 둔다).
BADGE_COLORS = {
    "Incompatible": ("FEF2F2", "B91C1C"),
    "Caution": ("FFFBEB", "92400E"),
    "Compatible": ("F0FDF4", "15803D"),
    "Abstain": ("F1F5F9", "475569"),
}


def _category_from_ko(word: str) -> str | None:
    """판정 한글 표기(부적합/주의/적합/판단 보류) -> category 키. 표 셀 값이 정확히
    이 단어들 중 하나일 때만 매치(배지로 바꿀지 판단하는 유일한 기준점)."""
    w = re.sub(r"\s+", "", word.strip())
    return {"부적합": "Incompatible", "주의": "Caution", "적합": "Compatible", "판단보류": "Abstain"}.get(w)

# 팔레트/폰트: ui-ux-pro-max "B2B Service / Data-Dense Dashboard" 추천 기반
# (navy 중심 단일 accent + Pretendard, 그라디언트·이모지 아이콘 등 AI 티 나는 요소는 제외).
CSS = """
<style>
@import url('https://cdn.jsdelivr.net/gh/orioncactus/pretendard@v1.3.9/dist/web/variable/pretendardvariable-dynamic-subset.css');
@import url('https://fonts.googleapis.com/css2?family=JetBrains+Mono:wght@400;500;600&display=swap');

:root {
  --mv-bg: #F8FAFC;
  --mv-surface: #FFFFFF;
  --mv-border: #E2E8F0;
  --mv-text: #0F172A;
  --mv-text-muted: #64748B;
  --mv-text-faint: #94A3B8;
  --mv-accent: #2563EB;
  --mv-accent-weak: #EFF4FF;
  --mv-shadow: 0 1px 2px rgba(15, 23, 42, .04), 0 1px 1px rgba(15, 23, 42, .03);
  --mv-shadow-md: 0 4px 12px rgba(15, 23, 42, .06), 0 1px 2px rgba(15, 23, 42, .04);
}

html, body, .stApp, button, input, select, textarea,
.stMarkdown, .stText, [data-testid="stMetricValue"], [data-testid="stWidgetLabel"] {
  font-family: "Pretendard Variable", Pretendard, -apple-system, BlinkMacSystemFont,
               "Segoe UI", Roboto, Helvetica, Arial, sans-serif !important;
}

.stApp { background: var(--mv-bg); }
.block-container { padding-top: 1.6rem; padding-bottom: 3rem; padding-left: 1.5rem; padding-right: 1.5rem; max-width: 1480px; }
#MainMenu { visibility: hidden; }
footer { visibility: hidden; }
header[data-testid="stHeader"] { background: transparent; }

.mv-header { display: flex; align-items: center; gap: .8rem; margin-bottom: .1rem; }
.mv-logo {
  width: 38px; height: 38px; border-radius: 9px; flex-shrink: 0;
  background: var(--mv-text);
  display: flex; align-items: center; justify-content: center;
  color: #fff; font-weight: 700; font-size: .78rem; letter-spacing: .01em;
}
.mv-title { font-size: 1.4rem; font-weight: 700; color: var(--mv-text); margin: 0; line-height: 1.3; letter-spacing: -.01em; }
.mv-subtitle { color: var(--mv-text-muted); font-size: .84rem; margin-top: .1rem; }
.mv-pipeline {
  font-family: "JetBrains Mono", SFMono-Regular, Consolas, monospace !important;
  font-size: .72rem; color: var(--mv-text-faint); letter-spacing: .01em; margin: .55rem 0 0 0;
}

hr.mv-divider { border: none; border-top: 1px solid var(--mv-border); margin: 1.3rem 0; }

div[data-testid="stVerticalBlockBorderWrapper"] {
  border-radius: 12px !important;
  border-color: var(--mv-border) !important;
  box-shadow: var(--mv-shadow) !important;
  background: var(--mv-surface) !important;
}

.mv-badge {
  display: inline-flex; align-items: center; gap: .45rem;
  padding: .3rem .75rem; border-radius: 6px;
  font-size: .82rem; font-weight: 600; letter-spacing: .01em;
}
.mv-badge .dot { width: 6px; height: 6px; border-radius: 50%; }
.mv-badge--incompatible { background: #FEF2F2; color: #B91C1C; }
.mv-badge--incompatible .dot { background: #DC2626; }
.mv-badge--caution { background: #FFFBEB; color: #92400E; }
.mv-badge--caution .dot { background: #D97706; }
.mv-badge--compatible { background: #F0FDF4; color: #15803D; }
.mv-badge--compatible .dot { background: #16A34A; }
.mv-badge--abstain { background: #F1F5F9; color: #475569; }
.mv-badge--abstain .dot { background: #64748B; }

.mv-grade-tag {
  font-family: "JetBrains Mono", SFMono-Regular, Consolas, monospace !important;
  font-size: .72rem; color: var(--mv-text-muted); background: var(--mv-bg);
  padding: .15rem .5rem; border-radius: 5px; border: 1px solid var(--mv-border);
}
.mv-mono {
  font-family: "JetBrains Mono", SFMono-Regular, Consolas, monospace !important;
  font-size: .8rem;
}

.mv-abstain-note {
  font-size: .82rem; color: #92400E; background: #FFFBEB; border: 1px solid #FDE68A;
  border-left: 3px solid #D97706;
  border-radius: 6px; padding: .5rem .8rem; margin-bottom: .4rem;
}
.mv-kr-panel {
  font-size: .88rem; color: var(--mv-text); background: var(--mv-accent-weak);
  border-left: 3px solid var(--mv-accent);
  border-radius: 6px; padding: .55rem .8rem; margin-bottom: .6rem;
}
.mv-disclaimer {
  border: 1px solid var(--mv-border); background: var(--mv-surface); color: var(--mv-text-muted);
  border-left: 3px solid var(--mv-text-faint);
  border-radius: 8px; padding: .85rem 1rem; font-size: .8rem; line-height: 1.6;
}
.mv-section-label {
  font-weight: 700; font-size: .82rem; color: var(--mv-text); margin: 1.6rem 0 .6rem 0;
  text-transform: uppercase; letter-spacing: .04em;
}

/* 물질 배치 판정 - <details> 아코디언 3개. 화면에서만 쓰고 DOCX/PDF 는 소제목+불릿이다. */
.mv-placement details {
  border: 1px solid var(--mv-border); border-radius: 8px; background: var(--mv-surface);
  padding: .55rem .85rem; margin-bottom: .45rem;
}
.mv-placement summary {
  cursor: pointer; font-weight: 700; font-size: .86rem; color: var(--mv-text);
  list-style: none;
}
.mv-placement summary::marker, .mv-placement summary::-webkit-details-marker { display: none; }
.mv-placement summary::before { content: "▸ "; color: var(--mv-text-faint); }
.mv-placement details[open] > summary::before { content: "▾ "; }
.mv-placement details[open] > summary { margin-bottom: .5rem; }
.mv-placement ul { margin: .2rem 0 .6rem 0; padding-left: 1.1rem; }
.mv-placement li { font-size: .85rem; line-height: 1.75; color: var(--mv-text); }
.mv-placement p { font-size: .85rem; line-height: 1.7; margin: .1rem 0 .3rem 0; }
.mv-place-sub {
  font-weight: 700; font-size: .78rem; color: var(--mv-text-muted);
  margin: .6rem 0 .1rem 0 !important; letter-spacing: .02em;
}

/* 선택한 물질쌍 패널 - 한글 풀이를 한 곳에 모았다(종전에는 좌/우로 나뉘어 있었다). */
.mv-pair-title {
  font-family: "JetBrains Mono", SFMono-Regular, Consolas, monospace !important;
  font-size: .82rem; color: var(--mv-text); margin: .55rem 0 .1rem 0;
}
.mv-pair-kr {
  background: var(--mv-accent-weak); border-left: 3px solid var(--mv-accent);
  border-radius: 6px; padding: .55rem .8rem; margin: .5rem 0 .6rem 0;
}
.mv-pair-group { font-weight: 700; font-size: .82rem; margin: .35rem 0 .1rem 0 !important; }
.mv-pair-kr p:first-child { margin-top: 0 !important; }
.mv-pair-kr ul { margin: 0 0 .1rem 0; padding-left: 1.05rem; }
.mv-pair-kr li { font-size: .8rem; line-height: 1.6; color: var(--mv-text-muted); }

/* FINAL REPORT 본문(LLM 마크다운)의 섹션 계층 - "## n. 제목" 헤더는 앱 전체에서
   이 보고서에만 나오므로 전역 선택자로 걸어도 다른 화면과 안 겹친다. */
[data-testid="stMarkdownContainer"] h2 {
  font-size: 1.08rem; font-weight: 700; color: var(--mv-text);
  margin: 1.7rem 0 .7rem 0; padding-bottom: .5rem;
  border-bottom: 2px solid var(--mv-text);
}
[data-testid="stMarkdownContainer"] h2:first-child { margin-top: 0; }
[data-testid="stMarkdownContainer"] h3 {
  font-size: .95rem; font-weight: 700; color: var(--mv-accent); margin: 1.1rem 0 .4rem 0;
}
[data-testid="stMarkdownContainer"] table {
  font-size: .86rem; width: 100%; table-layout: fixed; border-collapse: collapse; margin: .4rem 0 1rem 0;
}
[data-testid="stMarkdownContainer"] th {
  background: var(--mv-bg); text-align: left; padding: .5rem .7rem; border: 1px solid var(--mv-border);
}
[data-testid="stMarkdownContainer"] td {
  padding: .5rem .7rem; border: 1px solid var(--mv-border); word-break: break-word; vertical-align: top;
}
[data-testid="stMarkdownContainer"] blockquote {
  margin: .6rem 0 1.1rem 0; padding: .7rem 1rem; background: #FFFBEB; border: 1px solid #FDE68A;
  border-left: 3px solid #D97706; border-radius: 6px; color: #7C2D12; font-size: .86rem; line-height: 1.65;
}
[data-testid="stMarkdownContainer"] blockquote p { margin: 0; }

.mv-step {
  display: flex; align-items: center; gap: .55rem; margin: 1.9rem 0 .7rem 0;
}
.mv-step .num {
  width: 22px; height: 22px; border-radius: 50%; background: var(--mv-text); color: #fff;
  font-size: .72rem; font-weight: 700; display: flex; align-items: center; justify-content: center;
  flex-shrink: 0;
}
.mv-step .label { font-weight: 700; font-size: .92rem; color: var(--mv-text); }
.mv-step .hint { font-size: .78rem; color: var(--mv-text-muted); margin-left: .3rem; }

button[kind="primary"], .stButton > button[kind="primary"] {
  background: var(--mv-text) !important; border: none !important;
  box-shadow: none !important; font-weight: 600 !important;
}
button[kind="primary"]:hover { background: #1E293B !important; }

div[data-baseweb="select"] > div, .stMultiSelect [data-baseweb="select"] > div {
  border-radius: 8px !important; border-color: var(--mv-border) !important;
}

details[data-testid="stExpander"] {
  border: 1px solid var(--mv-border) !important; border-radius: 10px !important;
  margin-bottom: .5rem; box-shadow: none !important; background: var(--mv-surface) !important;
}
details[data-testid="stExpander"] summary { font-size: .83rem; font-weight: 600; color: var(--mv-text); }

[data-testid="stAlert"] { border-radius: 8px; font-size: .85rem; box-shadow: none; }

section[data-testid="stSidebar"] { background: var(--mv-surface); border-right: 1px solid var(--mv-border); }
.mv-side-label {
  font-size: .66rem; letter-spacing: .08em; text-transform: uppercase; color: var(--mv-text-faint);
  font-weight: 700; margin: 1.1rem 0 .4rem 0;
}
.mv-side-kv { display: flex; justify-content: space-between; font-size: .81rem; color: var(--mv-text-muted); padding: .2rem 0; }
.mv-side-kv span:last-child {
  font-family: "JetBrains Mono", SFMono-Regular, Consolas, monospace !important;
  color: var(--mv-text); font-size: .78rem;
}
</style>
"""


def _inject_css() -> None:
    st.markdown(CSS, unsafe_allow_html=True)


def step_header(n: int, label: str, hint: str = "") -> None:
    hint_html = f'<span class="hint">{hint}</span>' if hint else ""
    st.markdown(
        f'<div class="mv-step"><span class="num">{n}</span>'
        f'<span class="label">{label}</span>{hint_html}</div>',
        unsafe_allow_html=True,
    )


def badge_html(category: str, compact: bool = False) -> str:
    """compact=True: 한글 판정어만(부적합/주의/적합/판단 보류) - 보고서 표처럼 같은
    배지가 여러 행 반복되는 곳에서 " · Incompatible" 같은 영문 접미사가 매번 붙으면
    표가 장황해지므로 표 전용으로 축약형을 쓴다."""
    cls = category.lower()
    label = CATEGORY_LABEL.get(category, category)
    if compact:
        label = label.split(" · ")[0]
    return f'<span class="mv-badge mv-badge--{cls}"><span class="dot"></span>{label}</span>'


@st.cache_resource(show_spinner="검색 인덱스 로딩 중…")
def search_index():
    chunk_ids, texts, meta, vecs = [], [], [], []
    for tag in CORPUS_TAGS:
        full = R.load_corpus(GRAN, corpus_tag=tag)
        dvecs = R.embed_corpus(MODEL, GRAN, full, corpus_tag=tag)
        keep = [i for i, m in enumerate(full.meta) if m["section"] in SECTIONS]
        chunk_ids += [full.chunk_ids[i] for i in keep]
        texts += [full.texts[i] for i in keep]
        meta += [full.meta[i] for i in keep]
        vecs.append(dvecs[keep])
    corpus = R.Corpus(chunk_ids=chunk_ids, texts=texts, meta=meta)
    return (corpus, R.build_faiss(np.vstack(vecs)), R.build_bm25(BM25_TAG, corpus),
            R.boilerplate_penalty_vector(corpus))


@st.cache_resource(show_spinner="임베딩 모델 로딩 중…")
def embedder():
    R.torch_threads()
    from sentence_transformers import SentenceTransformer

    m = SentenceTransformer(R.EMBEDDING_MODELS[MODEL])
    m.max_seq_length = R.MAX_SEQ_LEN
    return m


@st.cache_data
def substances() -> dict[str, str]:
    """CAS -> 한글 물질명. 서비스 코퍼스(corpus_tag='service') 안에서 rag_chunks가
    쓰는 이름 그대로다 - 평가셋 질의문과 같은 이름을 유지하려고 registry 표준명으로
    덮지 않는다(표시명은 main()의 display_names가 따로 만든다)."""
    con = sqlite3.connect(DB_PATH)
    rows = con.execute(
        "select rc.cas_number, rc.chemical_name from rag_chunks rc "
        "join rag_corpus_membership m on m.cas_number = rc.cas_number "
        "where m.corpus_tag in ({}) group by rc.cas_number order by rc.chemical_name".format(
            ",".join("?" * len(CORPUS_TAGS))
        ),
        CORPUS_TAGS,
    ).fetchall()
    con.close()
    return dict(rows)


@st.cache_data
def registry() -> dict[str, dict]:
    """substance_registry 전체(CAS -> 한글명/영문명/기호/별칭). 물질 선정과 검색
    매칭용 식별 정보의 단일 기준(CORE 237종). 선택 목록은 여기서 KOSHA 미등재분만
    빼서 만든다 — "Registry ∪ 173" 규칙은 폐기됐다(docs/REGISTRY.md 5절).
    질의 텍스트(retrieve 쪽)와는 무관 — 그건 substances()가 그대로 담당한다
    (frozen eval 보존, build_substance_registry.py 참고)."""
    con = sqlite3.connect(DB_PATH)
    cols = ("cas_number", "name_ko", "name_en", "formula", "aliases")
    rows = con.execute(f"select {','.join(cols)} from substance_registry").fetchall()
    con.close()
    return {r[0]: dict(zip(cols[1:], r[1:])) for r in rows}


@st.cache_data
def kosha_info() -> dict[str, dict]:
    """CAS -> KOSHA MSDS 등재 정보(chemId / KOSHA 물질명 / 최종 갱신일).

    출처는 msds_chem_id_cache — kosha_msds_collector가 getChemList 결과를 적재해 둔
    기존 캐시다. 앱은 여기만 읽는다(런타임 API 호출 없음, 갱신은
    scripts/kosha_registry_lookup.py --fetch 담당). chem_id가 NULL인 행은
    "KOSHA에 미등재"로 확인된 물질이다."""
    con = sqlite3.connect(DB_PATH)
    rows = con.execute(
        "select cas_number, chem_id, chem_name_kor, last_date from msds_chem_id_cache"
    ).fetchall()
    con.close()
    return {r[0]: {"chem_id": r[1], "kosha_name": r[2], "last_date": r[3]} for r in rows}


@st.cache_data
def msds_detail(cas: str) -> pd.DataFrame:
    """CAS -> KOSHA MSDS 상세정보(§2 유해성·위험성 / §3 구성성분 / §9 물리화학 /
    §10 안정성·반응성). 출처는 msds_sections - kosha_msds_collector가 적재해 둔
    getChemDetail0X 응답이며, 앱은 읽기만 한다(런타임 API 호출 없음).

    표시명은 registry canonical name을 쓴다 - 여기서 KOSHA 원문명(§3 물질명)을
    가져다 라벨로 쓰지 않는다."""
    con = sqlite3.connect(DB_PATH)
    rows = con.execute(
        "select section, item_name_kor, lev, item_detail from msds_sections "
        "where cas_number=? order by section, ordr_idx",
        (cas,),
    ).fetchall()
    con.close()
    return pd.DataFrame(
        [
            {
                "구분": SECTION_LABEL.get(sec, f"§{sec}"),
                "항목": "  " * (max(lev or 1, 1) - 1) + (item or "-"),
                "내용": (detail or "").replace("|", chr(10)),
            }
            for sec, item, lev, detail in rows
            if detail
        ]
    )


@st.cache_data
def cameo_mapped() -> set[str]:
    """CAMEO 반응성 그룹이 실제로 붙은 CAS 집합. 그룹이 없으면 compatibility_engine이
    무조건 Abstain하므로(judge_pair_by_cas), 화면 안내를 정확히 나누는 데 쓴다.
    chemicals에 행만 있고 그룹이 없는 물질은 판정 불가이므로 여기 포함하지 않는다."""
    con = sqlite3.connect(DB_PATH)
    out = {c for (c,) in con.execute(
        "select distinct c.cas_number from chemicals c"
        " join chemical_group_membership m on m.chemical_id = c.chemical_id")}
    con.close()
    return out


# ── 근거 수집: 두 경로 ──────────────────────────────────────────────────────
# 현재 서비스는 사용자가 드롭다운에서 CAS를 고르는 closed-world 질의라, 근거가
# 결정론적으로 정해진다 -> msds_context() (production path. 쌍은 pair_context() 래퍼).
# retrieve()는 그 전제가 깨질 때를 위한 검색 계층 진입점이다. 지우지 않는 이유는
# 아래 docstring에 적었다. 두 함수는 역할이 다르며 서로를 호출하지 않는다.


def retrieve(query: str, topk: int = TOPK) -> list[dict]:
    """자유 텍스트 질의 -> §2·§10 하이브리드 검색 top-k. **현재 서비스 경로가 아니다.**

    근거 수집은 msds_context()가 처리한다. 이 함수를 남겨두는 건 다음 전제가 깨지는
    순간 바로 필요해지기 때문이다 — 셋 다 지금은 성립하지 않는다:
      - 입력이 CAS가 아니라 자유 문장이 되는 경우(물질을 문장에서 찾아내야 함)
      - §2·§10 밖의 섹션을 근거로 써야 하는 경우
      - 물질·청크가 늘어 한 물질 안에서도 골라야 하는 경우(현재는 물질당 2.1청크)

    죽은 코드가 아니다: `--check`의 질의 별칭 확장 테스트가 이 경로를 실제로 태워
    검증한다(페로센 질의가 자기 청크를 되찾는지). 검색이 깨지면 그 단정이 깨진다.

    측정된 검색 지표(docs/RETRIEVAL.md)는 이 계층의 것이고, run_ab.py /
    freeze_retrieval.py가 src/retrieval.py를 직접 써서 낸다 — 이 함수를 거치지 않는다.
    """
    corpus, index, bm25, penalty = search_index()
    qvec = embedder().encode([query], normalize_embeddings=True, convert_to_numpy=True).astype("float32")
    d = R.dense_rank(index, qvec, CAND_K)
    b = R.bm25_rank(bm25, [query], CAND_K)
    fused = R.rrf_fuse([d, b], topk, penalty=penalty)[0]
    return [
        {"chunk_id": corpus.chunk_ids[i], "text": corpus.texts[i], **corpus.meta[i]}
        for i in fused
        if i >= 0
    ]


def msds_context(cas_list: list[str]) -> list[dict]:
    """**서비스 기본 경로.** 주어진 물질들이 가진 §2·§10 청크 전부. 검색하지 않는다(2026-08-29).

    왜 검색을 안 쓰나: gold_evidence는 정의상 대상 물질의 §2다. 그리고 이 함수는 CAS를
    이미 인자로 받는다 - 검색기는 이미 결정론적으로 아는 것을 다시 찾고 있었다.
    코퍼스는 173종 371청크라 물질당 청크가 2개뿐이고, 그래서 top-k를 키우면 늘어나는
    건 전부 남의 물질이다.

    실측(service 2,240쌍질의 기준, 쌍당 평균):
      방식                  청크    정답확보   제3물질   프롬프트   질의임베딩
      분해 top-5/측         9.62    0.9888    60.1%    8,570자    461ms
      분해 top-3/측         5.93    0.9821    44.5%    5,326자    461ms
      CAS 직접조회(이 함수)  4.29    1.0000     0.0%    3,170자      0ms

    정렬을 고정한다 - 프롬프트의 [근거 n] 인용 번호가 이 순서에 대응한다.
    scripts/run_cameo_full.py의 pair_chunk_ids()와 같은 규칙이다(Generation 평가 경로).
    """
    order = {}
    for cas in cas_list:
        order.setdefault(cas, len(order))
    holes = ",".join("?" * len(order))
    con = sqlite3.connect(DB_PATH)
    rows = con.execute(
        "select c.chunk_id, c.text, c.cas_number, c.section, c.chemical_name from rag_chunks c "
        "join rag_corpus_membership m on m.cas_number = c.cas_number "
        f"where m.corpus_tag = '{SERVICE_CORPUS_TAG}' and c.granularity = 'section' "
        f"and c.section in (2, 10) and c.cas_number in ({holes})",
        tuple(order),
    ).fetchall()
    con.close()
    rows.sort(key=lambda r: (order.get(r[2], 9), r[3], r[0]))
    return [
        {"chunk_id": r[0], "text": r[1], "cas_number": r[2], "section": r[3], "chemical_name": r[4]}
        for r in rows
    ]


def pair_context(cas_a: str, cas_b: str) -> list[dict]:
    """쌍 전용 래퍼. 쌍 설명(explain)은 여기를 통해 들어온다."""
    return msds_context([cas_a, cas_b])


MAX_QUERY_ALIASES = 3  # 질의문(=LLM 프롬프트에도 들어감) 가독성 상한


@st.cache_data
def query_aliases() -> dict[str, list[str]]:
    """CAS -> 질의 확장용 별칭 목록.

    청크 헤더는 KOSHA 원문명으로 렌더돼 있어서, registry 표준명만으로 질의하면
    BM25가 어휘 매칭을 못 한다 - "페로센"으로 물어도 청크는 "디시클로펜타디에닐 철"
    이라 top-10에 자기 근거가 안 잡힌다. CAS로 묶인 다른 이름을 질의에 덧붙여 메운다.

    우선순위는 청크 헤더에 실제로 쓰인 이름(rag_chunks.chemical_name) > KOSHA 원문명
    > registry 영문명/별칭. 띄어쓰기만 다른 표기도 그대로 남긴다 - 형태소 토크나이저가
    두 표기를 같게 자른다는 보장이 없어서 둘 다 넣는 편이 안전하다.

    registry는 건드리지 않는다 - 전부 DB에 이미 있는 이름을 모으기만 한다."""
    con = sqlite3.connect(DB_PATH)
    out: dict[str, list[str]] = {}
    for cas, chunk_name, kosha_name, name_en, aliases, name_ko in con.execute(
        "select r.cas_number, "
        "       (select rc.chemical_name from rag_chunks rc where rc.cas_number=r.cas_number limit 1), "
        "       c.chem_name_kor, r.name_en, r.aliases, r.name_ko "
        "from substance_registry r left join msds_chem_id_cache c on c.cas_number = r.cas_number"
    ):
        # aliases는 공백 구분이지만 "염산 hydrochloric acid"처럼 여러 단어짜리가
        # 섞여 있어 쪼개지 않고 통째로 넣는다(BM25는 어차피 토크나이즈한다).
        cand = [chunk_name, kosha_name, name_en, aliases]
        seen, picked = {name_ko.lower()}, []
        for x in cand:
            x = (x or "").strip()
            if not x or x.lower() in seen:
                continue
            seen.add(x.lower())
            picked.append(x)
        out[cas] = picked[:MAX_QUERY_ALIASES]
    con.close()
    return out


def query_term(cas: str, name: str) -> str:
    """질의문에 쓸 물질 표기. 별칭이 있으면 "표준명(별칭, 별칭)"으로 붙인다."""
    extra = [a for a in query_aliases().get(cas, []) if a.lower() != name.lower()]
    return f"{name}({', '.join(extra)})" if extra else name


def explain(cas_a: str, cas_b: str, name_a: str, name_b: str) -> dict:
    """쌍 하나에 대한 전체 경로: CAMEO 조회 -> MSDS 근거 수집 -> LLM 설명.

    근거는 pair_context()로 두 CAS의 §2·§10을 직접 조회한다(검색 안 함, 사유는 거기 주석).
    질의문은 LLM에게 줄 질문 문장이라 query_term으로 별칭을 붙인 종전 형태 그대로 둔다.
    검색 경로(run_ab.py / freeze_retrieval.py)의 지표는 여기를 거치지 않는다."""
    query = QUERY_TEMPLATE.format(a=query_term(cas_a, name_a), b=query_term(cas_b, name_b))
    con = sqlite3.connect(DB_PATH)
    cameo = CL.lookup(con.cursor(), cas_a, cas_b)
    con.close()
    cameo_ctx = CL.format_context(cameo, name_a, name_b, detailed=True)
    contexts = pair_context(cas_a, cas_b)
    prompt = build_prompt(query, cameo_ctx, contexts)
    return {"query": query, "cameo": cameo, "cameo_ctx": cameo_ctx, "contexts": contexts, "prompt": prompt}


def render_matrix(verdict, names: dict[str, str]) -> None:
    """N x N 판정 매트릭스를 색상 히트맵 표로 렌더링 (verdict.to_table()의 시각화 버전)."""
    lookup: dict[tuple[str, str], str] = {}
    for (a, b), v in zip(itertools.combinations(verdict.inputs, 2), verdict.pair_verdicts):
        lookup[(a, b)] = lookup[(b, a)] = v.category

    labels = [names.get(c, c) for c in verdict.inputs]
    rows = []
    for a in verdict.inputs:
        rows.append(["—" if a == b else lookup[(a, b)] for b in verdict.inputs])
    df = pd.DataFrame(rows, index=labels, columns=labels)

    try:
        style_fn = df.style.map if hasattr(df.style, "map") else df.style.applymap
        styler = style_fn(lambda v: CELL_STYLE.get(v, "")).set_properties(
            **{"text-align": "center", "font-size": "0.8rem", "border-color": "#F0F1F3"}
        )
        st.dataframe(styler, use_container_width=True)
    except Exception:  # noqa: BLE001 - jinja2 등 스타일링 의존성 미설치 환경 대비 폴백
        st.dataframe(df, use_container_width=True)


def render_legend() -> None:
    st.markdown(
        " &nbsp; ".join(badge_html(cat) for cat in ("Incompatible", "Caution", "Compatible", "Abstain")),
        unsafe_allow_html=True,
    )


def verdict_rows(verdict, names: dict[str, str]) -> list[dict]:
    rows = []
    for (ca, cb), v in zip(itertools.combinations(verdict.inputs, 2), verdict.pair_verdicts):
        rows.append(
            {
                "물질A": names.get(ca, ca),
                "물질B": names.get(cb, cb),
                "판정": CATEGORY_LABEL.get(v.category, v.category),
                "근거등급": v.evidence_grade,
                "_category": v.category,
            }
        )
    return rows


def render_pair_detail(v, name_a: str, name_b: str) -> None:
    """선택한 쌍의 전부 - 판정 / 한글 풀이 / 근거 원문(접힘).

    종전에는 같은 쌍을 왼쪽(한글 풀이)과 오른쪽(영문 원인·근거)에 나눠 그렸다. 한 쌍을
    보려고 화면 양끝을 오가야 해서 한 곳으로 합쳤다. 영문 원문은 아래 아코디언에 남긴다.
    근거는 MSDS 가 아니라 CAMEO 68그룹 매트릭스이며, 쌍을 누를 때마다 검색·LLM 을 태우지
    않는 것은 의도된 선택이다."""
    st.markdown(
        badge_html(v.category)
        + f'<div class="mv-pair-title">{name_a} × {name_b}</div>',
        unsafe_allow_html=True,
    )
    for n in v.abstain_notes:
        st.markdown(f'<div class="mv-abstain-note">{n}</div>', unsafe_allow_html=True)

    if not v.group_pair_details:
        st.caption("참고할 CAMEO 그룹 근거가 없습니다.")
        return

    seen, blocks = set(), []
    for d in v.group_pair_details:
        key = (d.group_a_name, d.group_b_name)
        if key in seen:
            continue
        seen.add(key)
        hz = KRG.kr_codes(d.hazard_codes, KRG.HAZARD_KR)
        gp = KRG.kr_codes(d.gas_products, KRG.GAS_KR)
        rows = "".join(f"<li>{x}</li>" for x in hz)
        if gp:
            rows += "<li>발생 가능 물질: " + ", ".join(gp) + "</li>"
        blocks.append(
            f'<p class="mv-pair-group">{KRG.kr_group(d.group_a_name)} × '
            f'{KRG.kr_group(d.group_b_name)}</p>'
            + (f"<ul>{rows}</ul>" if rows else "")
        )
    st.markdown('<div class="mv-pair-kr">' + "".join(blocks) + "</div>", unsafe_allow_html=True)
    st.markdown(
        f'<span class="mv-grade-tag">CAMEO 반응성 그룹 매트릭스 · {v.evidence_grade}</span>',
        unsafe_allow_html=True,
    )

    with st.expander("근거 원문(영문) →"):
        for r in v.reasons:
            st.caption(r)
        for d in v.group_pair_details:
            st.markdown(f"**{d.group_a_name} × {d.group_b_name}** — {d.category}")
            extra = ", ".join(x for x in (
                d.description,
                f"hazard={d.hazard_codes}" if d.hazard_codes else None,
                f"gas={d.gas_products}" if d.gas_products else None,
            ) if x)
            if extra:
                st.caption(extra)
    st.caption("※ CAMEO 68그룹 매트릭스 용어를 그대로 옮긴 것으로, 단독 최종 판단 근거로 쓸 수 없습니다.")


# ---------- 관리자(ADMIN) 읽기 전용 대시보드 ----------
# 여기 값은 전부 freeze_retrieval.py / README 실측치에 맞춰 고정된 상수다(app 상단부 참고).
# 이 화면은 그 상수를 보여주기만 하고 바꾸지 않는다 - 바꾸면 실측치와 어긋난다.

def admin_tree() -> dict[str, dict[str, str]]:
    n_sub = len(substances())
    return {
        "Model": {
            "Generation Model": L.MODEL,
        },
        "Embedding": {
            "Embedding Model": MODEL,
        },
        "Retrieval": {
            "Section": "§" + " / §".join(map(str, sorted(SECTIONS))),
            "Top-K": str(TOPK),
            "Candidate-K": str(CAND_K),
            "Fusion": "Hybrid (Dense + BM25, RRF) + §10 penalty",
        },
        "Prompt": {
            "Generation Prompt": PROMPT_VERSION,
        },
        "Corpus": {
            "Corpus Version": " + ".join(CORPUS_TAGS),
            "Substance Count": str(n_sub),
        },
        "Pipeline": {
            "Retrieval": f"{MODEL} dense + BM25 hybrid, top-{TOPK}",
            "Generation": f"{L.MODEL} (reasoning_effort={GB.REASONING_EFFORT})",
            "Evaluation": "오프라인 배치 전용(eval_generation.py) — 이 화면엔 없음",
        },
    }


def render_admin_panel() -> None:
    with st.container(border=True):
        st.markdown('<div class="mv-section-label" style="margin-top:0;">관리자 · 파이프라인 설정 (읽기 전용)</div>', unsafe_allow_html=True)
        st.caption("아래 값은 README 실측치에 맞춰 고정된 상수입니다. 이 화면에서 바꿀 수 없습니다.")
        cols = st.columns(3)
        for i, (section, kv) in enumerate(admin_tree().items()):
            with cols[i % 3]:
                st.markdown(f'<div class="mv-side-label">{section}</div>', unsafe_allow_html=True)
                for k, v in kv.items():
                    st.markdown(
                        f'<div class="mv-side-kv"><span>{k}</span><span>{v}</span></div>',
                        unsafe_allow_html=True,
                    )


# ---------- 최종 종합 보고서 (신규 LLM 호출 1회, 쌍별 재검색 없음) ----------
# v2(2026-08-29): "부적합 목록 + 정성 서술" 에서 "관리 그룹 중심" 으로 재편했다.
# 코드가 결정론적으로 확정하는 것은 셋뿐이다 - 관리 그룹 구성, 판정, 근거 번호.
# 문장은 전부 LLM 이 쓴다. 숫자(건수/비율)는 미리 계산해 "그대로 옮길 것"으로 못박아
# LLM 이 통계를 잘못 계산하는 걸 막는다.

# 5절이었던 고정 문구. LLM 에게 복사시키면 복사를 틀릴 기회만 생기므로 코드가 붙인다
# (화면 + DOCX + PDF 세 경로 공통).
JUDGEMENT_BASIS = (
    "본 판정은 CAMEO 68그룹 반응성 매트릭스 대조 결과([참고자료] 등급)에 기반하며, "
    "단독으로 최종 위험성평가의 근거로 사용할 수 없다. 실제 취급·저장 판단에는 KOSHA "
    "MSDS 원문과 산업안전보건법상 요구사항, 전문가 검토를 반드시 병행해야 한다."
)

# 표 셀에 넣을 짧은 위험 라벨. kr_glossary.HAZARD_KR 은 "쉬운 설명" 패널용 완전한 문장이라
# 표에 넣으면 같은 문장이 행마다 반복돼 보고서가 장황해진다(실측). 강도 표현("~일 수 있음")은
# 범례가 아니라 판정이 이미 담고 있으므로 여기서는 명사구로만 쓴다.
HAZARD_SHORT = {
    "C": "부식성 생성물", "E": "폭발성 생성물", "F": "인화성 생성물",
    "G": "기체 발생·압력 상승", "NR": "알려진 위험 반응 없음", "R1": "발열",
    "R2": "가열 시 불안정", "R3": "격렬·폭발적 반응", "R4": "격렬한 중합",
    "T": "독성 생성물", "UR": "위험 가능(근거 불충분)",
}

FINAL_REPORT_INSTRUCTION = """
역할: 화학물질을 어떻게 나눠 보관·관리할지 보여주는 짧은 보고서에서 **두 개 절만** 쓴다.
배치 판정과 데이터 제약은 시스템이 뒤에 직접 붙이므로 쓰지 않는다.

# 입력 데이터 계약
- 판정(부적합/주의/적합/판단 보류)은 코드가 이미 확정했다. 다시 계산하거나 바꾸지 않는다.
- [MSDS 근거]에는 번호가 붙어 있다. 본문에 원문을 옮기지 말고 근거 번호만 [근거 3] 형태로
  붙인다. 원문과 출처는 화면의 별도 근거 영역이 보여준다.
- 값이 없는 항목은 "확인되지 않음"으로 쓰고 추정하지 않는다.

# 내용 규칙
1. 제공된 CAMEO 근거와 MSDS 근거에 있는 것만 쓴다. 일반 화학지식으로 보완하지 않는다.
2. 근거보다 강하거나 약하게 표현하지 않는다.
   - may / can / possible -> 가능성("~할 수 있다")
   - should / recommended -> 권고
   - must / prohibited / incompatible -> 필수·금지
3. 판정은 CAMEO 그룹쌍 대조 결과다. 실제 반응의 발생·강도·결과를 단정하지 않으며 이는
   위험 방향과 안전 방향 모두에 적용된다.
   - 금지(위험 단정): "폭발한다", "화재 위험이 크다", "~가 발생한다"
   - 금지(안전 단정): "실제 위험성은 제한적이다", "우려할 수준은 아니다"
4. 개별 물질의 유해성과 물질 간 조합 판정을 혼동하지 않는다. 물질별 정보를 서로 섞지 않는다.
5. 반응 메커니즘 추론과 화학적 추론 과정을 쓰지 않는다. 근거에 없는 조치(감시·설비·농도
   조건 등)를 새로 만들지 않는다.
6. 반복 서술, 일반적인 안전수칙, 앞 절에 이미 나온 내용의 재나열을 하지 않는다.

# 서식 규칙
7. 각 표는 제시된 헤더 행과 구분선 행(|---|)을 문자 그대로 복사한 뒤 데이터 행만 채운다.
   컬럼 개수·순서 변경과 헤더 병합을 금지한다.
8. 중첩 bullet과 번호 목록을 쓰지 않는다. 표 셀은 한 줄로 짧게 유지하고, 셀 안에서
   줄을 바꾸지 않는다 - 줄을 바꾸면 표가 깨진다.
9. 판정 단어는 "부적합"/"주의"/"적합"/"판단 보류" 중 정확히 하나만 쓴다(화면에서 배지로
   자동 변환됨).
10. 표기 언어는 한국어. 물질명은 주어진 표기를 그대로 쓴다 - MSDS 근거 안에 다른 이름
    (원문명·이명)이 보여도 그 이름으로 바꾸지 않는다.
11. 이 지시문의 문장이나 지시 어투("~라고 쓴다", "~하지 않는다")가 본문에 나타나면 안 된다.

아래 2개 섹션 제목을 순서·문구 그대로 "## n. 제목" 형식으로 쓴다. 다른 절을 만들지 않는다.

## 1. 개별 물질 유의사항

물질별로 한 줄씩, 그 물질에 배정된 근거에서 확인되는 핵심 위험·특성만 쓴다. **항목은 3개
이내**로 추리고 관리 판단에 영향을 주지 않는 정보는 싣지 않는다. 물질은 하나도 빠뜨리지
않는다 - 근거가 없는 물질도 행을 만들되 "확인되지 않음"으로 쓴다.
"근거" 열에는 "물질별 근거 번호"에 주어진 그 물질의 표기를 그대로 옮긴다 - 다른 물질의
번호를 쓰지 않는다.

| 물질 | 핵심 유의사항 | 근거 |
|---|---|---|

## 2. 분리 필요 관계

주어진 "분리 필요 관계" 표를 헤더부터 마지막 행까지 그대로 옮긴다. 행 순서와 셀 값을 바꾸지
않고 행을 빼거나 합치지 않으며, 해석·귀결을 덧붙이지 않는다. 표 외에 다른 문장을 쓰지 않는다.
자세한 CAMEO 그룹쌍·코드 원문은 화면 오른쪽 물질쌍 패널이 보여주므로 여기에 옮기지 않는다.

분리가 필요한 조합이 없으면 표 대신 "분리가 필요한 조합은 없다." 한 줄만 쓴다.

"""


def compatible_groups(verdict) -> list[list[str]]:
    """서로 전부 Compatible 인 물질끼리 묶는다 - 이 묶음이 "함께 관리 가능한 그룹"이다.

    Caution/Incompatible/Abstain 은 전부 "함께 묶지 않음"으로 취급한다. Caution 을
    Compatible 과 같게 다루면 그룹 구성 자체가 판정을 완화하는 셈이 된다.

    LLM 에게 시키지 않는 이유는 이게 판정의 연장이기 때문이다 - 어떤 물질을 한 칸에
    넣어도 되는가는 CAMEO 판정이 결정하지 설명자가 결정하지 않는다.

    ponytail: greedy clique cover - 그룹 수가 최소라는 보장은 없다. 최소 분할이
    필요해지면 그때 정확 알고리즘으로 교체한다(N<=20이라 완전탐색도 가능하다).
    """
    ok = {}
    for (a, b), v in zip(itertools.combinations(verdict.inputs, 2), verdict.pair_verdicts):
        ok[(a, b)] = ok[(b, a)] = v.category == "Compatible"
    groups: list[list[str]] = []
    for cas in verdict.inputs:
        for g in groups:
            if all(ok[(cas, other)] for other in g):
                g.append(cas)
                break
        else:
            groups.append([cas])
    return groups


def _major_risk_pairs(verdict, names: dict[str, str]) -> list[dict]:
    """부적합/주의 쌍만 골라 2절 표의 셀 값을 그대로 만든다.

    CAMEO 그룹쌍과 코드 원문은 넣지 않는다 - 화면 오른쪽 물질쌍 패널이 같은 내용을
    보여주고 있어서 표에 다시 실으면 보고서만 길어진다."""
    out = []
    for (ca, cb), v in zip(itertools.combinations(verdict.inputs, 2), verdict.pair_verdicts):
        if v.category not in ("Incompatible", "Caution"):
            continue
        worst = [d for d in v.group_pair_details if d.category == v.category]
        d0 = worst[0] if worst else None
        codes = [c.strip() for c in ((d0.hazard_codes if d0 else None) or "").split(",") if c.strip()]
        risk = " · ".join(HAZARD_SHORT.get(c, c) for c in codes)
        if not risk:
            # 코드가 없는 쌍(그룹 자기조합 등)은 설명이 유일한 근거다.
            risk = (((d0.description if d0 and d0.description else "")
                     or "; ".join(v.reasons))[:80] or "확인되지 않음")
        out.append({
            "판정": CATEGORY_LABEL.get(v.category, v.category).split(" · ")[0],
            "물질A": names.get(ca, ca), "물질B": names.get(cb, cb),
            "위험": risk,
        })
    # 정렬해서 준다 - LLM 에게 정렬을 시키면 정렬을 틀릴 기회를 준다.
    rank = {"부적합": 0, "주의": 1}
    out.sort(key=lambda m: (rank.get(m["판정"], 9), -len(m["위험"])))
    return out


def build_final_report_prompt(verdict, rows: list[dict], names: dict[str, str],
                              contexts: list[dict]) -> str:
    lines = [f"입력 물질 {len(verdict.inputs)}종 조합 판정", ""]

    # 근거 번호 -> 물질. 어느 번호가 누구 것인지는 코드가 안다.
    cite = {}
    for n, c in enumerate(contexts, 1):
        cite.setdefault(c["cas_number"], []).append(str(n))

    lines.append("물질별 근거 번호(1절에서 그 물질 행에만 이 번호를 쓴다):")
    for cas in verdict.inputs:
        ns = cite.get(cas)
        lines.append(f"- {names.get(cas, cas)}"
                     + (f" [근거 {', '.join(ns)}]" if ns else " [근거 없음]"))

    major = _major_risk_pairs(verdict, names)
    # 표를 통째로 만들어 준다 - key=value 로 주면 LLM 이 "코드=" 같은 라벨까지 셀에
    # 복사하거나 이웃 행의 값을 가져온다(둘 다 실측됨).
    lines += ["", "분리 필요 관계(부적합/주의 전건, 정렬 완료):"]
    if major:
        cols = ["판정", "물질A", "물질B", "위험"]
        lines.append("| 판정 | 물질 A | 물질 B | 주요 위험 |")
        lines.append("|---|---|---|---|")
        for m in major:
            lines.append("| " + " | ".join(m[c].replace("|", "/") for c in cols) + " |")
    else:
        lines.append("없음(분리가 필요한 조합 없음)")

    # 근거 번호 규약은 run_cameo_context_pilot.build_prompt 와 같다([근거 n] + chunk_id).
    # 물질명·섹션을 같은 줄에 덧붙인다 - 20종이면 누구의 청크인지가 번호만으로는 안 보인다.
    lines += ["", "[MSDS 근거] (본문에는 번호만 인용하고 원문을 옮기지 않는다)"]
    for n, c in enumerate(contexts, 1):
        label = names.get(c["cas_number"], c["chemical_name"])
        sec = SECTION_LABEL.get(c["section"], "§" + str(c["section"]))
        lines.append(f"\n[근거 {n}] (chunk_id={c['chunk_id']}) {label} {sec}\n{c['text']}")

    lines += ["", FINAL_REPORT_INSTRUCTION]
    return "\n".join(lines)


# ── 물질 배치 판정 (전부 코드 산출) ─────────────────────────────────────────
# 판정과 관리 그룹에서 기계적으로 따라오는 결과라 LLM 을 태울 이유가 없다.
# 화면은 <details> 아코디언으로, DOCX/PDF 는 소제목+불릿 마크다운으로 같은 데이터를 낸다.

PLACEMENT_RULES = [
    "부적합 조합 → 분리 보관",
    "주의 조합 → 동일 구역 배치 여부 검토",
    "단독 관리가 필요한 물질 → 별도 배치",
]
PLACEMENT_LIMIT = (
    "본 결과는 CAMEO에 등록된 반응성 정보를 기준으로 산출되었습니다. "
    "CAMEO에 반영되지 않은 추가 반응 가능성은 포함하지 않으며, 새로운 데이터가 확보되면 "
    "재평가가 필요합니다."
)


def _group_label(cas_list: list[str], verdict, limit: int = 3) -> str:
    """물질들이 속한 CAMEO 반응성 그룹의 한글명을 이어 붙인 라벨."""
    prof = {pr.cas: pr for pr in verdict.profiles}
    out: list[str] = []
    for cas in cas_list:
        for g in (prof[cas].groups if cas in prof else []):
            kr = KRG.kr_group(g)
            if kr not in out:
                out.append(kr)
    return " · ".join(out[:limit]) or "분류 정보 없음"


def placement_plan(verdict, contexts: list[dict], names: dict[str, str]) -> dict:
    """배치 결과 = 관리 그룹(동일 구역/단독) + 부적합 쌍(분리) + 주의 쌍(검토)."""
    groups = compatible_groups(verdict)
    nm = lambda c: names.get(c, c)  # noqa: E731

    together = [(_group_label(g, verdict), " + ".join(nm(c) for c in g))
                for g in groups if len(g) > 1]
    alone = [(_group_label(g, verdict), nm(g[0])) for g in groups if len(g) == 1]

    separate, review = [], []
    for (ca, cb), v in zip(itertools.combinations(verdict.inputs, 2), verdict.pair_verdicts):
        if v.category not in ("Incompatible", "Caution"):
            continue
        item = (f"{_group_label([ca], verdict, 2)} ↔ {_group_label([cb], verdict, 2)}",
                f"{nm(ca)} ↔ {nm(cb)}")
        (separate if v.category == "Incompatible" else review).append(item)

    limits = []
    unmapped = [pr.cas for pr in verdict.profiles if not pr.mapped]
    if unmapped:
        limits.append("CAMEO 반응성 그룹 매핑이 없어 판정하지 못한 물질: "
                      + ", ".join(nm(c) for c in unmapped))
    have = {c["cas_number"] for c in contexts}
    no_ev = [c for c in verdict.inputs if c not in have]
    if no_ev:
        limits.append("MSDS §2·§10 근거가 없는 물질: " + ", ".join(nm(c) for c in no_ev))
    n_abstain = sum(1 for v in verdict.pair_verdicts if v.category == "Abstain")
    if n_abstain:
        limits.append(f"근거 부족으로 판정하지 않은(판단 보류) 조합 {n_abstain}건")
    limits.append("농도·온도·상(phase)·혼합비 정보가 입력에 없어 반응의 강도와 발생 "
                  "가능성은 이 결과만으로 확정할 수 없습니다")

    return {"together": together, "alone": alone, "separate": separate,
            "review": review, "limits": limits}


_PLACEMENT_BLOCKS = (
    ("동일 구역 배치", "together"),
    ("단독 관리", "alone"),
    ("분리 배치", "separate"),
    ("배치 검토 (주의)", "review"),
)


def placement_markdown(plan: dict) -> str:
    """DOCX/PDF 용. 화면과 같은 내용을 소제목+불릿으로 낸다."""
    out = ["## 3. 물질 배치 판정", "", "### 배치 원칙", ""]
    out += [f"- {r}" for r in PLACEMENT_RULES]
    out += ["", "### 배치 결과", ""]
    for title, key in _PLACEMENT_BLOCKS:
        if plan[key]:
            out.append(f"**{title}**")
            out += [f"- {label}: {members}" for label, members in plan[key]]
            out.append("")
    out += ["### 판정의 한계", "", PLACEMENT_LIMIT, ""]
    out += [f"- {x}" for x in plan["limits"]]
    return "\n".join(out)


def placement_html(plan: dict) -> str:
    """화면용 <details> 아코디언. 기본은 접힌 상태이고 배치 결과만 펼쳐 둔다."""
    def block(title, items):
        rows = "".join(f"<li><b>{label}</b>: {members}</li>" for label, members in items)
        return f"<p class='mv-place-sub'>{title}</p><ul>{rows}</ul>"

    body = "".join(block(t, plan[k]) for t, k in _PLACEMENT_BLOCKS if plan[k])
    rules = "".join(f"<li>{r}</li>" for r in PLACEMENT_RULES)
    limits = "".join(f"<li>{x}</li>" for x in plan["limits"])
    return (
        "<div class='mv-placement'>"
        "<details><summary>배치 원칙</summary>"
        f"<ul>{rules}</ul></details>"
        "<details open><summary>배치 결과</summary>"
        f"{body}</details>"
        "<details><summary>판정의 한계</summary>"
        f"<p>{PLACEMENT_LIMIT}</p><ul>{limits}</ul></details>"
        "</div>"
    )


def generate_final_report(verdict, rows: list[dict], names: dict[str, str]) -> dict:
    contexts = msds_context(list(verdict.inputs))
    prompt = build_final_report_prompt(verdict, rows, names, contexts)
    plan = placement_plan(verdict, contexts, names)
    try:
        data = L.chat(
            [{"role": "user", "content": prompt}],
            max_tokens=GB.MAX_TOKENS,
            reasoning_effort=GB.REASONING_EFFORT,
        )
        content, err = data["choices"][0]["message"]["content"].rstrip(), None
    except Exception as e:  # noqa: BLE001 - 데모에서 스택트레이스 대신 사유만
        content, err = None, f"{type(e).__name__}: {str(e)[:300]}"
    # contexts 를 같이 돌려준다 - 본문의 [근거 n] 을 화면에서 원문으로 펼치려면
    # 프롬프트에 넣은 것과 **같은 순서의 같은 리스트**여야 한다.
    return {"content": content, "error": err, "contexts": contexts, "plan": plan,
            "prompt_chars": len(prompt)}


# ---------- DOCX / PDF 내보내기 ----------

def _split_report_sections(report_text: str) -> list[tuple[str, str]]:
    """"## n. 제목" 헤더 기준으로 (제목, 본문) 리스트로 쪼갠다. 헤더가 하나도 없으면
    전체를 본문 하나로 취급(LLM이 지시 형식을 안 지킨 경우 대비 폴백)."""
    parts = re.split(r"(?m)^#{1,3}\s*\d*\.?\s*(.+)$", report_text)
    if len(parts) == 1:
        return [("최종 분석 결과", report_text.strip())]
    out = []
    for i in range(1, len(parts), 2):
        title, body = parts[i].strip(), parts[i + 1].strip() if i + 1 < len(parts) else ""
        out.append((title, body))
    return out


_TABLE_ROW_RE = re.compile(r"^\s*\|")
_QUOTE_ROW_RE = re.compile(r"^\s*>")
_BULLET_ROW_RE = re.compile(r"^\s*[-*]\s+")
_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_BADGE_CELL_RE = re.compile(r"(\|\s*)\*{0,2}([가-힣\s]{2,6})\*{0,2}(\s*)(?=\|)")


def _parse_table(lines: list[str]) -> tuple[list[str], list[list[str]]]:
    """마크다운 파이프 테이블(헤더/구분선/데이터행)을 (헤더, 행 리스트)로 파싱."""
    def cells(line: str) -> list[str]:
        line = line.strip().strip("|")
        return [c.strip() for c in line.split("|")]

    header = cells(lines[0])
    body = [cells(l) for l in lines[2:] if l.strip()]  # lines[1] = --- 구분선
    return header, body


def _parse_blocks(body: str) -> list[tuple[str, object]]:
    """섹션 본문을 (table|quote|bullet|para, data) 블록 리스트로 분해한다.
    DOCX/PDF export와 Streamlit 배지 렌더링이 이 파서 하나를 공유해서 세
    출력이 항상 같은 구조로 나오게 한다."""
    lines = body.split("\n")
    blocks: list[tuple[str, object]] = []
    i, n = 0, len(lines)
    while i < n:
        line = lines[i]
        if not line.strip():
            i += 1
        elif _TABLE_ROW_RE.match(line):
            j = i
            while j < n and _TABLE_ROW_RE.match(lines[j]):
                j += 1
            if j - i >= 2:
                blocks.append(("table", _parse_table(lines[i:j])))
            i = j
        elif _QUOTE_ROW_RE.match(line):
            j, quote = i, []
            while j < n and _QUOTE_ROW_RE.match(lines[j]):
                quote.append(_QUOTE_ROW_RE.sub("", lines[j]).strip())
                j += 1
            blocks.append(("quote", " ".join(t for t in quote if t)))
            i = j
        elif _BULLET_ROW_RE.match(line):
            blocks.append(("bullet", _BULLET_ROW_RE.sub("", line).strip()))
            i += 1
        else:
            j, para = i, []
            while j < n and lines[j].strip() and not (
                _TABLE_ROW_RE.match(lines[j]) or _QUOTE_ROW_RE.match(lines[j]) or _BULLET_ROW_RE.match(lines[j])
            ):
                para.append(lines[j].strip())
                j += 1
            blocks.append(("para", " ".join(para)))
            i = j
    return blocks


def _bold_runs(text: str) -> list[tuple[str, bool]]:
    """"**굵게**" 마크다운을 (텍스트, bold여부) 런 리스트로 쪼갠다 - DOCX/PDF는
    HTML을 못 쓰니 직접 폰트 굵기를 적용해야 KPI 수치 강조가 살아난다."""
    runs, pos = [], 0
    for m in _BOLD_RE.finditer(text):
        if m.start() > pos:
            runs.append((text[pos:m.start()], False))
        runs.append((m.group(1), True))
        pos = m.end()
    if pos < len(text):
        runs.append((text[pos:], False))
    return runs or [("", False)]


def _badgify_markdown(text: str) -> str:
    """표 셀 값이 정확히 판정 단어(부적합/주의/적합/판단 보류)일 때만 배지 HTML로
    치환한다. 파이프를 소비하지 않고 lookahead로만 확인해서 인접 셀 파싱이
    깨지지 않게 한다."""
    def repl(m: re.Match) -> str:
        cat = _category_from_ko(m.group(2))
        return f"{m.group(1)}{badge_html(cat, compact=True)}{m.group(3)}" if cat else m.group(0)
    return _BADGE_CELL_RE.sub(repl, text)


def _docx_add_runs(paragraph, text: str) -> None:
    for chunk, bold in _bold_runs(text):
        if chunk:
            paragraph.add_run(chunk).bold = bold


def _docx_shade_cell(cell, hex_color: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _docx_add_table(doc, header: list[str], rows: list[list[str]]) -> None:
    from docx.shared import RGBColor

    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light List Accent 1"
    for cell, text in zip(table.rows[0].cells, header):
        cell.paragraphs[0].add_run(text).bold = True
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            p = cell.paragraphs[0]
            cat = _category_from_ko(text)
            if cat:
                bg, fg = BADGE_COLORS[cat]
                _docx_shade_cell(cell, bg)
                run = p.add_run(text)
                run.bold = True
                run.font.color.rgb = RGBColor.from_string(fg)
            else:
                _docx_add_runs(p, text)


def _docx_add_callout(doc, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    _docx_shade_cell(cell, "FFFBEB")
    run = cell.paragraphs[0].add_run("⚠ 데이터 제약 사항 — " + text)
    run.italic = True


def export_text(report: dict) -> str:
    """DOCX/PDF 용 본문. 화면은 배치 판정을 <details> 로 그리지만 문서는 소제목+불릿이다."""
    return report["content"] + "\n\n" + placement_markdown(report["plan"])


def build_docx_bytes(report_text: str, rows: list[dict]) -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading("SMART-MSDS 최종 분석 결과", level=1)
    for title, body in _split_report_sections(report_text):
        doc.add_heading(title, level=2)
        for kind, data in _parse_blocks(body):
            if kind == "table":
                header, trows = data
                if trows:
                    _docx_add_table(doc, header, trows)
            elif kind == "quote" and data:
                _docx_add_callout(doc, data)
            elif kind == "bullet":
                _docx_add_runs(doc.add_paragraph(style="List Bullet"), data)
            elif kind == "para" and data:
                _docx_add_runs(doc.add_paragraph(), data)

    doc.add_heading("판정 기준 및 참고사항", level=2)
    doc.add_paragraph(JUDGEMENT_BASIS)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _pdf_bold_markup(text: str) -> str:
    text = text.replace("&", "&amp;").replace("<", "&lt;")
    return _BOLD_RE.sub(r"<b>\1</b>", text)


_PDF_TABLE_WEIGHTS = {2: (0.35, 0.65), 3: (0.5, 1.3, 2.2), 5: (0.9, 1.3, 1.3, 1.6, 3.0)}


def _pdf_col_widths(header: list[str], total: float) -> list[float]:
    weights = _PDF_TABLE_WEIGHTS.get(len(header), tuple([1.0] * len(header)))
    s = sum(weights)
    return [total * w / s for w in weights]


def _pdf_table(header: list[str], rows: list[list[str]], font: str):
    from reportlab.lib import colors
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import Paragraph, Table, TableStyle

    hstyle = ParagraphStyle("th", fontName=font, fontSize=8.5, leading=11, textColor=colors.HexColor("#0F172A"))
    cstyle = ParagraphStyle("td", fontName=font, fontSize=8.5, leading=12)

    style_cmds = [
        ("FONTNAME", (0, 0), (-1, -1), font),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#E2E8F0")),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F1F5F9")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]
    data = [[Paragraph(h, hstyle) for h in header]]
    for r, row in enumerate(rows, start=1):
        cells = []
        for c, text in enumerate(row):
            cat = _category_from_ko(text)
            if cat:
                bg, fg = BADGE_COLORS[cat]
                style_cmds.append(("BACKGROUND", (c, r), (c, r), colors.HexColor(f"#{bg}")))
                pstyle = ParagraphStyle(f"badge{r}_{c}", parent=cstyle, textColor=colors.HexColor(f"#{fg}"))
                cells.append(Paragraph(f"<b>{text}</b>", pstyle))
            else:
                cells.append(Paragraph(_pdf_bold_markup(text), cstyle))
        data.append(cells)

    tbl = Table(data, colWidths=_pdf_col_widths(header, 450), hAlign="LEFT")
    tbl.setStyle(TableStyle(style_cmds))
    return tbl


def build_pdf_bytes(report_text: str, rows: list[dict]) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    FONT = "HYSMyeongJo-Medium"  # reportlab 내장 CID 폰트 - 별도 폰트파일 없이 한글 지원
    pdfmetrics.registerFont(UnicodeCIDFont(FONT))
    h1 = ParagraphStyle("h1", fontName=FONT, fontSize=16, leading=20, spaceAfter=12)
    h2 = ParagraphStyle("h2", fontName=FONT, fontSize=12, leading=16, spaceBefore=10, spaceAfter=6)
    body_style = ParagraphStyle("body", fontName=FONT, fontSize=9.5, leading=14)
    bullet_style = ParagraphStyle("bullet", parent=body_style, leftIndent=12, bulletIndent=2)
    quote_style = ParagraphStyle(
        "quote", fontName=FONT, fontSize=9, leading=13, textColor=colors.HexColor("#7C2D12"),
        backColor=colors.HexColor("#FFFBEB"), borderPadding=8, leftIndent=4, spaceBefore=4, spaceAfter=4,
    )

    story = [Paragraph("SMART-MSDS 최종 분석 결과", h1)]
    for title, sect_body in _split_report_sections(report_text):
        story.append(Paragraph(title, h2))
        for kind, data in _parse_blocks(sect_body):
            if kind == "table":
                header, trows = data
                if trows:
                    story.append(_pdf_table(header, trows, FONT))
                    story.append(Spacer(1, 6))
            elif kind == "quote" and data:
                story.append(Paragraph("⚠ 데이터 제약 사항 — " + _pdf_bold_markup(data), quote_style))
            elif kind == "bullet":
                story.append(Paragraph("• " + _pdf_bold_markup(data), bullet_style))
            elif kind == "para" and data:
                story.append(Paragraph(_pdf_bold_markup(data), body_style))

    story.append(Paragraph("판정 기준 및 참고사항", h2))
    story.append(Paragraph(JUDGEMENT_BASIS, body_style))

    buf = io.BytesIO()
    SimpleDocTemplate(buf, pagesize=A4).build(story)
    return buf.getvalue()


def main() -> None:
    st.set_page_config(
        page_title="MSDS 혼재보관 위험성 판정",
        page_icon="🧪",
        layout="wide",
        initial_sidebar_state="collapsed",
    )
    _inject_css()

    head_l, head_r = st.columns([5, 1.4])
    with head_l:
        st.markdown(
            '<div class="mv-header"><div class="mv-logo">SM</div>'
            '<div><p class="mv-title">SMART-MSDS · 물질 조합 안전성 분석</p>'
            '</div></div>',
            unsafe_allow_html=True,
        )
    with head_r:
        st.markdown('<div style="height:.4rem;"></div>', unsafe_allow_html=True)
        b1, b2 = st.columns(2)
        with b1:
            if st.button("관리자", use_container_width=True):
                st.session_state["show_admin"] = not st.session_state.get("show_admin", False)
        with b2:
            db_ok = DB_PATH.exists()
            st.markdown(
                f'<span class="mv-badge mv-badge--{"compatible" if db_ok else "incompatible"}">'
                f'<span class="dot"></span>{"정상" if db_ok else "DB 없음"}</span>',
                unsafe_allow_html=True,
            )
    st.markdown('<hr class="mv-divider"/>', unsafe_allow_html=True)

    if st.session_state.get("show_admin"):
        render_admin_panel()

    names = substances()  # RAG 질의문 기준 이름(frozen eval 보존, 절대 변경 금지)
    reg = registry()
    kosha = kosha_info()
    # 표시용 이름의 기준은 substance_registry의 canonical name이다 - 선택 목록
    # (_label)·상세정보·판정 결과·보고서가 전부 같은 이름을 쓰게 하기 위함.
    # registry에 없는 코퍼스 전용 물질만 rag_chunks 이름(names)으로 폴백한다.
    # names는 RAG 질의문 생성(explain)에 그대로 쓰이므로 frozen eval은 무영향.
    display_names = {**names, **{cas: r["name_ko"] for cas, r in reg.items()}}

    def _label(cas: str) -> str:
        """Streamlit multiselect는 라벨 문자열 자체에 substring 필터를 건다(커스텀
        검색 위젯 불필요). 한글/영문/기호/별칭/CAS를 전부 라벨에 넣어 무엇을
        입력하든 매칭되게 한다 — "아연"도 "zinc"도 "Zn"도 같은 물질로 찾힘."""
        r = reg.get(cas)
        if not r:
            return f"{names.get(cas, cas)} ({cas})"
        parts = [r["name_ko"]]
        extra = [x for x in (r["name_en"], r["formula"], r["aliases"]) if x]
        if extra:
            parts.append("(" + ", ".join(extra) + ")")
        parts.append(f"[{cas}]")
        return " ".join(parts)

    # 서비스 물질 선정의 기준은 substance_registry(CORE 5축) 하나다 — 2026-08-22 확정
    # 237종. 과거의 173 코퍼스나 "Registry ∪ 173" 규칙은 선정 기준이 아니며, 코퍼스는
    # 검색 인덱스·평가 재현용 자산으로만 남는다(rag_corpus_membership).
    # 여기서 추가로 걸러내는 건 KOSHA MSDS 미등재분뿐이다 — 상세정보를 줄 수 없는
    # 물질을 고르게 하지 않기 위함이고, Registry 237종에서 빼는 게 아니다.
    # 미등재 판정 근거: msds_chem_id_cache.chem_id IS NULL. getChemList를 CAS(searchCnd=1)/
    # 국문명·영문명(searchCnd=0) 3경로로 실조회해 39종 전부 0건 확인
    # (data/collection/kosha_unlisted_39.csv, results/kosha_missing39_probe_2026-08-22.csv).
    known_cas = set(reg)
    unlisted = {cas for cas in known_cas if not (kosha.get(cas) or {}).get("chem_id")}
    all_cas = known_cas - unlisted
    labels = {_label(cas): cas for cas in all_cas}
    no_evidence = {cas for cas in all_cas if cas not in names}

    col_a, col_b, col_c = st.columns([1, 2.5, 1.2], gap="medium")

    with col_a:
        step_header(1, "ANALYSIS", "물질 선택")
        with st.container(border=True):
            picked = st.multiselect(
                "물질 선택",
                sorted(labels),
                max_selections=20,
                label_visibility="collapsed",
                placeholder="물질명을 검색해 2종 이상 선택하세요 (예: 에탄올, 질산 …)",
            )
            st.caption(f"2~20종 선택 가능 · 현재 {len(picked)}종 · 선택 가능 {len(all_cas)}종")
            analyze = st.button("분석 시작", type="primary", disabled=len(picked) < 2, use_container_width=True)

    if analyze:
        st.session_state["analyzed_picked"] = picked
        st.session_state.pop("final_report_cache", None)
        st.session_state.pop("selected_pair", None)
    analyzed = st.session_state.get("analyzed_picked")

    if not analyzed or len(analyzed) < 2 or not set(analyzed) <= set(labels):
        with col_b:
            st.info("왼쪽에서 물질을 2종 이상 선택하고 '분석 시작'을 누르면 최종 보고서가 표시됩니다.")
        with col_c:
            st.caption("분석된 물질쌍이 없습니다.")
        st.markdown('<hr class="mv-divider"/>', unsafe_allow_html=True)
        st.markdown(f'<div class="mv-disclaimer">{DISCLAIMER}</div>', unsafe_allow_html=True)
        return

    cas_list = [labels[p] for p in analyzed]
    eng = CompatibilityEngine(str(DB_PATH))
    verdict = eng.judge_combination_by_cas(cas_list)
    eng.close()

    # 근거 없음 안내는 두 갈래다 — CAMEO 매핑이 있으면 판정은 되고 설명 근거만 없지만,
    # 매핑까지 없으면 판정 자체가 Abstain이다. 한 문구로 묶으면 후자를 오도한다.
    mapped = cameo_mapped()
    no_ev = [c for c in cas_list if c in no_evidence]
    if [c for c in no_ev if c in mapped]:
        st.caption(
            "검색 근거(MSDS §2/§10 청크) 없음: "
            + ", ".join(display_names.get(c, c) for c in no_ev if c in mapped)
            + " — CAMEO 반응성 그룹 판정은 반영되나 원문 근거는 붙지 않습니다."
        )
    if [c for c in cas_list if c not in mapped]:
        st.caption(
            "CAMEO 반응성 그룹 매핑 없음: "
            + ", ".join(display_names.get(c, c) for c in cas_list if c not in mapped)
            + " — 관련 조합은 Abstain(판정 보류)으로 처리됩니다."
        )

    with st.expander("선택 물질의 KOSHA MSDS 상세정보"):
        st.dataframe(
            pd.DataFrame(
                [
                    {
                        "물질": display_names.get(c, c),
                        "CAS": c,
                        "KOSHA 물질명": (kosha.get(c) or {}).get("kosha_name") or "-",
                        "chemId": (kosha.get(c) or {}).get("chem_id") or "-",
                        "MSDS 최종 갱신": (kosha.get(c) or {}).get("last_date") or "-",
                    }
                    for c in cas_list
                ]
            ),
            hide_index=True,
            use_container_width=True,
        )
        st.caption(
            "출처: KOSHA MSDS Open API(getChemList) 조회 결과 캐시. 표시명은 "
            "substance_registry의 표준명을 쓰고 KOSHA 원문명은 대조용으로만 병기한다 — "
            "갱신은 scripts/kosha_registry_lookup.py --fetch"
        )

        detail_target = st.selectbox(
            "상세정보를 볼 물질",
            cas_list,
            format_func=lambda c: display_names.get(c, c),
            key="msds_detail_target",
        )
        detail = msds_detail(detail_target)
        st.markdown(
            f'<div class="mv-section-label">{display_names.get(detail_target, detail_target)} '
            f'· MSDS 상세</div>',
            unsafe_allow_html=True,
        )
        if detail.empty:
            st.caption(
                "KOSHA MSDS 상세정보가 없습니다(미등재 물질). "
                "판정은 CAMEO 반응성 그룹 근거로만 이뤄집니다."
            )
        else:
            for label in detail["구분"].unique():
                st.caption(label)
                st.dataframe(
                    detail.loc[detail["구분"] == label, ["항목", "내용"]],
                    hide_index=True,
                    use_container_width=True,
                    column_config={"내용": st.column_config.TextColumn(width="large")},
                )
            st.caption(
                "출처: KOSHA MSDS Open API getChemDetail02/03/09/10 — "
                "갱신은 scripts/kosha_msds_collector.py "
                "--target-csv data/collection/registry_core207.csv"
            )

    rows = verdict_rows(verdict, display_names)
    pairs = list(itertools.combinations(verdict.inputs, 2))
    fmt = lambda p: f"{display_names.get(p[0], p[0])} × {display_names.get(p[1], p[1])}"  # noqa: E731
    worst_idx = max(range(len(pairs)), key=lambda i: verdict.pair_verdicts[i].category == verdict.category)

    with col_b:
        step_header(2, "FINAL REPORT", "최종 결과 보고서")
        with st.container(border=True):
            st.markdown(badge_html(verdict.category), unsafe_allow_html=True)
            cache_key = tuple(cas_list)
            report_cache = st.session_state.setdefault("final_report_cache", {})
            cached_report = report_cache.get(cache_key)

            gen = st.button(
                "다시 생성" if cached_report else "최종 보고서 생성", type="primary", use_container_width=True
            )
            if gen:
                with st.spinner("최종 보고서 생성 중… (30초 내외)"):
                    report_cache[cache_key] = cached_report = generate_final_report(verdict, rows, display_names)

            if cached_report:
                if cached_report["error"]:
                    st.error(f"LLM 호출 실패: {cached_report['error']}")
                else:
                    st.markdown('<div class="mv-section-label" style="margin-top:1rem;">SMART-MSDS 최종 분석 결과</div>', unsafe_allow_html=True)
                    st.markdown(_badgify_markdown(cached_report["content"]), unsafe_allow_html=True)
                    st.markdown('<div class="mv-section-label">물질 배치 판정 완료</div>'
                                + placement_html(cached_report["plan"]),
                                unsafe_allow_html=True)
                    st.markdown(
                        '<div class="mv-section-label">판정 기준 및 참고사항</div>'
                        f'<div class="mv-disclaimer">{JUDGEMENT_BASIS}</div>',
                        unsafe_allow_html=True,
                    )
                    # 본문은 근거 번호만 인용한다. 원문·출처는 여기서 펼친다 -
                    # 번호는 프롬프트에 넣은 contexts 순서 그대로다. 물질별로 접어 둔다
                    # (20종이면 40청크라 한 덩어리로 펼치면 읽을 수가 없다).
                    ctxs = cached_report["contexts"]
                    by_cas: dict[str, list] = {}
                    for n, c in enumerate(ctxs, 1):
                        by_cas.setdefault(c["cas_number"], []).append((n, c))
                    st.markdown('<div class="mv-section-label">근거 원문 · MSDS '
                                f'{len(ctxs)}건</div>', unsafe_allow_html=True)
                    for cas, items in by_cas.items():
                        nums = ", ".join(str(n) for n, _ in items)
                        label = display_names.get(cas, items[0][1]["chemical_name"])
                        with st.expander(f"{label} · [근거 {nums}]"):
                            for n, c in items:
                                sec = SECTION_LABEL.get(c["section"], "§" + str(c["section"]))
                                st.caption(f"[근거 {n}] {sec} · {c['chunk_id']}")
                                st.text(c["text"])

                    dl1, dl2, dl3 = st.columns(3)
                    with dl1:
                        st.download_button(
                            "매핑 CSV 다운로드",
                            pd.DataFrame(rows).drop(columns="_category").to_csv(index=False).encode("utf-8-sig"),
                            file_name="msds_verdict_matrix.csv",
                            mime="text/csv",
                            use_container_width=True,
                        )
                    with dl2:
                        st.download_button(
                            "DOCX 다운로드",
                            build_docx_bytes(export_text(cached_report), rows),
                            file_name="msds_final_report.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True,
                        )
                    with dl3:
                        st.download_button(
                            "PDF 다운로드",
                            build_pdf_bytes(export_text(cached_report), rows),
                            file_name="msds_final_report.pdf",
                            mime="application/pdf",
                            use_container_width=True,
                        )
            else:
                st.caption("'최종 보고서 생성'을 누르면 물질별 유의사항 · 분리 필요 관계 · "
                           "물질 배치 판정으로 이뤄진 보고서가 생성됩니다.")

            if len(verdict.inputs) > 2:
                with st.expander("전체 반응 매트릭스 보기"):
                    render_legend()
                    render_matrix(verdict, display_names)

    with col_c:
        step_header(3, "PAIR INSPECTOR", "분석된 물질쌍")
        with st.container(border=True):
            st.markdown('<div class="mv-section-label" style="margin-top:0;">분석된 물질쌍</div>', unsafe_allow_html=True)
            selected = st.session_state.get("selected_pair")
            if selected not in pairs:
                selected = pairs[worst_idx]
            for i, (ca, cb) in enumerate(pairs):
                v = verdict.pair_verdicts[i]
                label = f"{SEVERITY_DOT.get(v.category, '⚪')} {fmt((ca, cb))}"
                if st.button(label, key=f"pair_{ca}_{cb}", use_container_width=True,
                             type="primary" if (ca, cb) == selected else "secondary"):
                    selected = (ca, cb)
                    st.session_state["selected_pair"] = selected

            st.markdown('<hr class="mv-divider"/>', unsafe_allow_html=True)
            st.markdown('<div class="mv-section-label" style="margin-top:0;">선택한 물질쌍</div>', unsafe_allow_html=True)
            sel_idx = pairs.index(selected)
            render_pair_detail(
                verdict.pair_verdicts[sel_idx],
                display_names.get(selected[0], selected[0]),
                display_names.get(selected[1], selected[1]),
            )

    st.markdown('<hr class="mv-divider"/>', unsafe_allow_html=True)
    st.markdown(f'<div class="mv-disclaimer">{DISCLAIMER}</div>', unsafe_allow_html=True)


if __name__ == "__main__":
    if "--check" in sys.argv:
        # LLM 없이 검색+CAMEO 경로만 점검(네트워크/API 키 불필요).
        n = substances()
        cas_a, cas_b = "7646-85-7", "7697-37-2"  # frozen 평가셋 t0 케이스
        got = explain(cas_a, cas_b, n[cas_a], n[cas_b])
        assert got["cameo"].category == "Incompatible", got["cameo"].category
        # 근거는 검색이 아니라 CAS 직접조회다(pair_context). 개수 대신 구성을 본다 -
        # 제3물질이 섞이지 않을 것, 양쪽 §2(=gold_evidence)가 반드시 들어있을 것.
        ids = [c["chunk_id"] for c in got["contexts"]]
        assert ids, "근거가 비어있음"
        assert all(c["cas_number"] in (cas_a, cas_b) for c in got["contexts"]),             f"제3물질 청크 혼입: {ids}"
        assert {f"sec::{cas_a}::2", f"sec::{cas_b}::2"} <= set(ids), f"§2 누락: {ids}"
        assert ids == sorted(ids, key=lambda x: ({cas_a: 0, cas_b: 1}[x.split("::")[1]],
                                                 int(x.split("::")[2].split("::")[0]), x)),             f"인용 번호가 붙는 순서가 고정되지 않음: {ids}"
        assert "[CAMEO 반응성 그룹 조회" in got["prompt"], "프롬프트에 CAMEO 컨텍스트 누락"
        assert n[cas_a] in got["query"], "질의문이 frozen 이름을 쓰지 않음"
        print("OK:", got["cameo"].category, len(ids), "chunks (제3물질 0, 양쪽 §2 포함)")

        # 표시명 일관성: registry에 있는 물질은 모든 화면에서 canonical name 하나만 쓴다.
        # (main()의 display_names와 동일한 식 - 여기서 깨지면 화면에서도 깨진다)
        reg = registry()
        display = {**n, **{c: r["name_ko"] for c, r in reg.items()}}
        assert all(display[c] == r["name_ko"] for c, r in reg.items()), "registry 표시명 미적용"
        assert display[cas_a] != n[cas_a], "rag_chunks 이름이 그대로 남아있음(뒤집기 실패)"

        # KOSHA 상세정보는 조회만 하고 표시명을 덮어쓰지 않는다(§3 물질명은 자료로만).
        d = msds_detail(cas_a)
        assert set(d["구분"]) == set(SECTION_LABEL.values()), sorted(set(d["구분"]))
        kosha_name = (kosha_info().get(cas_a) or {}).get("kosha_name")
        assert display[cas_a] == reg[cas_a]["name_ko"] != kosha_name, "KOSHA 원문명이 표시명을 덮어씀"
        print("OK: 표시명", display[cas_a], "| KOSHA 원문명", kosha_name, "| 상세", len(d), "행")

        # 검색 대상 필터: KOSHA 미등재 물질은 선택 목록에서 빠지되 registry에는 남는다.
        k = kosha_info()
        known = set(reg)   # 선정 기준은 Registry 단독 (docs/REGISTRY.md 4절)
        unlisted = {c for c in known if not (k.get(c) or {}).get("chem_id")}
        assert unlisted, "미등재 집합이 비어있음(캐시 미적재?)"
        assert all(msds_detail(c).empty for c in list(unlisted)[:3]), "미등재인데 상세정보가 있음"
        # Registry 규모도 숫자를 박지 않는다 - VIEW의 in_registry 와 일치하는지만 본다.
        _con = sqlite3.connect(DB_PATH)
        n_reg_view = _con.execute(
            "select count(*) from substance_status where in_registry=1").fetchone()[0]
        _con.close()
        assert len(reg) == n_reg_view, f"Registry 조회와 VIEW 불일치: {len(reg)} vs {n_reg_view}"
        assert set(reg) - unlisted, "registry가 통째로 걸러짐"
        print("OK: 검색대상", len(known - unlisted), "종 / 미등재 제외", len(unlisted),
              "종 (registry 잔존", len(set(reg) - unlisted), "종)")

        # 질의 별칭 확장: 청크 헤더가 KOSHA 원문명이라 표준명만으론 BM25가 못 잡는
        # 케이스(페로센 vs 디시클로펜타디에닐 철)가 실제로 복구되는지 확인한다.
        ferro = "102-54-5"
        term = query_term(ferro, reg[ferro]["name_ko"])
        assert "디시클로펜타디에닐" in term, term
        plain = retrieve(QUERY_TEMPLATE.format(a=reg[ferro]["name_ko"], b="수산화나트륨"))
        expand = retrieve(QUERY_TEMPLATE.format(a=term, b="수산화나트륨"))
        assert ferro not in {h["cas_number"] for h in plain}, "전제 붕괴: 확장 없이도 잡힘"
        assert ferro in {h["cas_number"] for h in expand}, "별칭 확장이 자기 청크를 못 살림"
        print("OK: 질의 별칭 확장 -", term)

        # 코퍼스 규모: 서비스 검색 대상은 corpus_tag='service' 하나다.
        # 개수를 여기에 적지 않는다 - substance_status VIEW가 계산한 값과 대조만 한다.
        # (숫자를 박으면 물질이 하나 늘 때마다 코드가 어긋난다. docs/REGISTRY.md 1절)
        con = sqlite3.connect(DB_PATH)
        tags = dict(con.execute(
            "select corpus_tag, count(*) from rag_corpus_membership"
            " where corpus_tag in ('173','core','service') group by corpus_tag"))
        expect_service = con.execute(
            "select count(*) from substance_status"
            " where service_eligible=1 and chunks_ready=1").fetchone()[0]
        gap = con.execute(
            "select count(*) from substance_status where index_status='인덱스 결손'").fetchone()[0]
        con.close()
        assert tags.get("service") == expect_service, \
            f"service 태그가 VIEW와 불일치: 태그 {tags.get('service')} vs VIEW {expect_service}"
        assert len(substances()) == expect_service, \
            f"인덱스 대상이 service와 불일치: {len(substances())} vs {expect_service}"
        assert gap == 0, f"인덱스 결손 {gap}종 - 서비스 자격이 있는데 청크가 없다"
        assert tags.get("173") and tags.get("core"), f"평가 재현용 태그 소실: {tags}"
        print("OK: service 코퍼스", len(substances()), "종 (인덱스 결손 0) |",
              "보존 태그 173:", tags["173"], "core:", tags["core"])

        # CAMEO 매핑: 그룹이 붙은 물질만 판정 가능하고, 없으면 무조건 Abstain이다.
        served = set(reg) - unlisted
        mapped_served = served & cameo_mapped()
        assert len(mapped_served) == expect_service, \
            f"CAMEO 매핑 물질과 service 코퍼스 불일치: {len(mapped_served)} vs {expect_service}"
        eng = CompatibilityEngine(str(DB_PATH))
        pair_mapped = eng.judge_pair_by_cas("7664-93-9", "1310-73-2")   # 황산 x 수산화나트륨
        assert pair_mapped.category != "Abstain", pair_mapped.category
        unmapped_cas = next(iter(served - mapped_served))
        pair_unmapped = eng.judge_pair_by_cas("7664-93-9", unmapped_cas)
        assert pair_unmapped.category == "Abstain", pair_unmapped.category
        print("OK: CAMEO 매핑", len(mapped_served), "/", len(served), "종 |",
              "매핑쌍", pair_mapped.category, "| 미매핑쌍", pair_unmapped.category)

        # 신규 매핑(pubchem_cameo_2026-08-22)이 실제로 판정으로 이어지는지 — 메탄올 x 질산.
        methanol = eng.judge_pair_by_cas("67-56-1", "7697-37-2")
        assert methanol.category == "Incompatible", methanol.category
        assert not msds_detail("67-56-1").empty, "메탄올 MSDS 상세 없음"
        print("OK: 신규 매핑 판정 - 메탄올 x 질산", methanol.category,
              "| 상세", len(msds_detail("67-56-1")), "행")

        # --- 최종 보고서 경로 ---
        # 코드가 결정론적으로 확정하는 셋만 검사한다: 관리 그룹 구성 / 근거 번호 순서 /
        # 판정 기준 고정 문구의 위치. 문장은 LLM 이 쓰므로 여기서 검사하지 않는다.
        combo_cas = sorted(mapped_served)[:6]
        ctxs = msds_context(combo_cas)
        cids = [c["chunk_id"] for c in ctxs]
        assert all(c["cas_number"] in combo_cas for c in ctxs), f"제3물질 혼입: {cids}"
        # §2 는 물질에 따라 sec::{cas}::2::p1 처럼 쪼개져 있다 - chunk_id 문자열이 아니라
        # (cas, section) 으로 확인한다.
        have = {(c["cas_number"], c["section"]) for c in ctxs}
        assert {(c, 2) for c in combo_cas} <= have, f"§2 누락: {sorted(have)}"
        assert cids == [c["chunk_id"] for c in msds_context(combo_cas)], "근거 순서가 비결정적"
        assert msds_context([cas_a, cas_b]) == pair_context(cas_a, cas_b), "pair_context 래퍼 불일치"

        cv = eng.judge_combination_by_cas(combo_cas)
        groups = compatible_groups(cv)
        pair_cat = {}
        for (pa, pb), pv in zip(itertools.combinations(cv.inputs, 2), cv.pair_verdicts):
            pair_cat[(pa, pb)] = pair_cat[(pb, pa)] = pv.category
        assert sorted(c for g in groups for c in g) == sorted(cv.inputs), f"물질 누락/중복: {groups}"
        for g in groups:
            for pa, pb in itertools.combinations(g, 2):
                assert pair_cat[(pa, pb)] == "Compatible", f"같은 그룹인데 {pair_cat[(pa, pb)]}"
        if any(c == "Compatible" for c in pair_cat.values()):
            assert any(len(g) > 1 for g in groups), "Compatible 쌍이 있는데 그룹이 전부 단독이다"
        print("OK: 관리 그룹", [[display.get(c, c) for c in g] for g in groups])

        rp = build_final_report_prompt(cv, verdict_rows(cv, display), display, ctxs)
        assert JUDGEMENT_BASIS not in rp, "고정 문구가 프롬프트에 남아있다"
        assert JUDGEMENT_BASIS not in FINAL_REPORT_INSTRUCTION, "고정 문구가 지시문에 남아있다"
        assert "[근거 1]" in rp and "물질별 근거 번호" in rp, "근거 데이터가 프롬프트에 없다"

        # 배치 판정은 전부 코드 산출이다 - 물질이 빠지거나 겹치면 안 된다.
        plan = placement_plan(cv, ctxs, display)
        placed = [m for _, members in plan["together"] + plan["alone"]
                  for m in members.split(" + ")]
        assert sorted(placed) == sorted(display.get(c, c) for c in cv.inputs),             f"배치에서 물질이 누락/중복: {placed}"
        md = placement_markdown(plan)
        for h in ("배치 원칙", "배치 결과", "판정의 한계"):
            assert h in md, h
        assert placement_html(plan).count("<details") == 3

        # export 는 LLM 없이 돈다 - 관리 그룹 3열 표가 새로 생겼으므로 그 경로를 태운다.
        sample = "\n".join([
            "## 1. 관리 그룹", "",
            "| 그룹 | 물질 | 공통 취급 유의사항 |", "|---|---|---|",
            "| 그룹 1 | 아세톤 | 확인되지 않음 |", "",
        ])
        assert len(build_docx_bytes(sample, [])) > 5000, "DOCX 생성 실패"
        assert len(build_pdf_bytes(sample, [])) > 1000, "PDF 생성 실패"
        print("OK: 보고서 프롬프트", len(rp), "자 /", len(combo_cas), "종 /",
              len(ctxs), "청크 / export 정상")
    else:
        main()
